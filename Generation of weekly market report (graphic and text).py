#!/usr/bin/env python
# coding: utf-8

# In[1]:
import os
import json
import docx
import xlrd #excel 讲道理这个是用来读取excel的，并且xlsx应该是没问题的，但是莫名不行
import xlwt #excel
import openpyxl #这个可以实现读取
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH #设置对象居中、对齐等对齐选项。
from docx.enum.text import WD_TAB_ALIGNMENT,WD_TAB_LEADER  #设置制表符等
from docx.shared import Inches,Cm   #设置图像大小
from docx.shared import Pt    #设置像素、缩进等
from docx.shared import RGBColor    #设置字体颜色
from docx.shared import Length    #设置宽度
from PIL import Image
from docx.oxml.ns import qn #设置中文字体
from datetime import datetime
import pandas as pd
import excel2img
# path = 'project_path'

PATH_1 = r'\\192.168.1.120\share\op_pro\OP_2024-08-19'
# PATH_1为周报素材的路径，可以根据日期需要手动修改，或者用下面这个%的now.时间
# PATH_1 = r'\\192.168.1.120\share\op_pro\OP_%s-%s-%s' % (now.year,now.month,now.day)

doc = docx.Document( PATH_1 + r"\素材\base.docx")
# 后面写入内容之后再doc=Document(PATH)
# 否则直接doc = Document(PATH)会报错，因为是空白文档

def word_create(name):
#     path= r"\\192.168.1.120\share\op_pro\OP_2024-05-13\\" #这个应该是后期要修改的，看素材文件夹放到哪里
    path = r".\\output_doc\\"
    path += name
    return path

def doc_create():
    now=datetime.now()
    name = "op_test_%s-%s-%s.docx" % (now.year,now.month,now.day)
    word_create(name)
    path = word_create(name)
    return path

# In[2]:
def file_import():
    #以下三个路径按需更改，不涉及到日期变化，分别是原先文件夹中的两个description的json
    file = PATH_1 + r"\素材\op.json"
    with open(file, 'r') as f:
        data = json.load(f)
    file_sentiment = PATH_1 + r"\素材\op_sentiment.json"
    with open(file_sentiment, 'r') as f:
        data_sentiment = json.load(f)
    file_description = PATH_1 + r"\素材\op_description.xlsx"

    # 添加固定文本，此路径需要更改
    file_fixed_txt = open(PATH_1 + r"\素材\fixed.txt", encoding='utf-8')
    fixed_txt = {}
    for line in file_fixed_txt.readlines():
        key = line.split('#')[0]
        value = line.split('#')[1]
        value = value.strip('\n')
        fixed_txt[key] = value
    return [data,data_sentiment,file_description,fixed_txt]


# In[3]:
# #整体字体设置
def doc_style():
    doc.styles['Normal'].font.name = 'Times New Roman' #设置西文字体
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体') #设置中文字体

# In[4]:
# 添加段落
def add_p():
    global p
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(21) #两个五号字体10.5 * 2 = 21磅
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示

#添加文字块
def add_r(key,doc,**filename):
    r = p.add_run(filename.get(key,"缺少指定内容"))
    r.font.size = Pt(10.5)
    doc.save(PATH)
    
#添加只有一个文字块的段落
def add_pg(key,doc,**filename):
    global p
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(21) #两个五号字体10.5 * 2 = 21磅
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    r = p.add_run(filename.get(key,"缺少指定内容"))
    r.font.size = Pt(10.5)
    doc.save(PATH)
    
#添加description文字块
def add_d(key,doc,filename):
    global p
    op_des = openpyxl.load_workbook(filename) #打开excel文件
    sheet = op_des.worksheets[0]
    rows = sheet.max_row + 1
    for i in range(2,rows):
        if key == sheet.cell(i,1).value:
            r = p.add_run(sheet.cell(i,2).value)
            r.font.size = Pt(10.5)
            r.font.bold = True
        else:
            i += 1
    doc.save(PATH)

#添加description，只针对第一部分的“一、”文字格式
def add_h1(h1,key,doc,filename):
    op_des = openpyxl.load_workbook(filename) #打开excel文件
    sheet = op_des.worksheets[0]
    rows = sheet.max_row + 1
    for i in range(2,rows):
        if key == sheet.cell(i,1).value:
            r1 = h1.add_run(sheet.cell(i,2).value)
            r1.font.color.rgb = RGBColor(31,73,125) # 蓝色
            r1.font.bold = True
            r1.font.size = Pt(14) #四号字体对应14磅
        else:
            i += 1
    doc.save(PATH)
    
# 添加图片
def add_pic(pic_name, **kwargs):
    global p
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0) #图片不缩进
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    r = p.add_run("")
    r.add_picture(pic_name,width=Cm(14.6), **kwargs)
    doc.save(PATH)
    
def add_excel_pic(pic_name):
    global p
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0) #图片不缩进
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    r = p.add_run("")
    r.add_picture(pic_name)
    doc.save(PATH)

#注意区分paragraph和run的属性，如果遇到换行问题，考虑在add_run和add_paragraph之间替换
# paragraph默认换行

#在同一行添加并列的两张图片
def add_pic2(p,pic_name):
    r = p.add_run("")
    r.add_picture(pic_name,width=Cm(7.3))
    doc.save(PATH)

#通过op_description的excel获取日期
def add_date(key,doc,filename):
    op_des = openpyxl.load_workbook(filename) #打开文件
    sheet = op_des.worksheets[0]
    for i in range(2,28):
        if key == sheet.cell(i,1).value:
            r = p.add_run(str(sheet.cell(i,2).value.strftime("%Y-%m-%d")))
            r.font.size = Pt(10.5)
            r.font.bold = True
        else:
            i += 1
    doc.save(PATH)

#写入日期
def date_period():
    p.add_run(r'（')
    p.add_run(add_date('上周第一个交易日',doc,file_description))
    p.add_run('至')
    p.add_run(add_date('上周最后一个交易日',doc,file_description))
    p.add_run(r'）：')
    
#从excel读入交集日期
def read_union_stock_market():
    union_date = openpyxl.load_workbook(PATH_1 + "\相似日期\交集日期.xlsx")
    sheet = union_date.worksheets[0]
    rows = sheet.max_row + 1
    list = []
    for i in range(2, rows):
        cell = sheet.cell(row=i, column=2)
        list.append(cell.value)
        i += 1
    if list == []:
        doc.add_paragraph("无更新；")
    else:
        doc.add_paragraph(str(list))
    doc.save(PATH)
    
def read_union_commodity_market():
    union_date = openpyxl.load_workbook(PATH_1 + "\相似日期\交集日期.xlsx")
    sheet = union_date.worksheets[1]
    rows = sheet.max_row + 1
    list = []
    for i in range(2, rows):
        cell = sheet.cell(row=i, column=2)
        list.append(cell.value)
        i += 1
    if list == []:
        doc.add_paragraph("无更新；")
    else:
        doc.add_paragraph(str(list))
    doc.save(PATH)
    
def read_union_industries():
    union_date = openpyxl.load_workbook(PATH_1 + "\相似日期\交集日期.xlsx")
    sheet = union_date.worksheets[2]
    rows = sheet.max_row + 1
    list = []
    for i in range(2, rows):
        cell = sheet.cell(row=i, column=2)
        list.append(cell.value)
        i += 1
    if list == []:
        doc.add_paragraph("无更新；")
    else:
        doc.add_paragraph(str(list))
    doc.save(PATH)
    
def read_union_factor_return():
    union_date = openpyxl.load_workbook(PATH_1 + "\相似日期\交集日期.xlsx")
    sheet = union_date.worksheets[3]
    rows = sheet.max_row + 1
    list = []
    for i in range(2, rows):
        cell = sheet.cell(row=i, column=2)
        list.append(cell.value)
        i += 1
    if list == []:
        doc.add_paragraph("无更新；")
    else:
        doc.add_paragraph(str(list))
    doc.save(PATH)

def path_pic():
    xl = pd.ExcelFile(PATH_1 + r"\素材\op_insert_pics.xlsx")
    sheet_name = xl.sheet_names
    path_pic = [r"\\192.168.1.120\share\op_pro\自动生成文件\\pic\\" + j + ".png"for j in sheet_name]
    return path_pic,sheet_name
    
def excel_pic(i,path_pic,sheet_name):
    excel_file = pd.read_excel(PATH_1 + r"\素材\op_insert_pics.xlsx", sheet_name = sheet_name[i])
    excel2img.export_img(PATH_1 + r"\素材\op_insert_pics.xlsx", path_pic[i], _range = '%s!A1:%s%s' % \
                                        (sheet_name[i],chr(96 + excel_file.shape[1]).upper(), excel_file.shape[0] + 1))
 
# In[5]:
def Head1():
    H1 = doc.add_paragraph()
    H1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    H1.paragraph_format.line_spacing = 1.5 # 1.5倍行距 单倍行距 1.0
    now=datetime.now()
    R1 = H1.add_run("Optimus Prime（%s年%s月%s日）"%(now.year,now.month,now.day))
    R1.font.bold = True
    R1.font.size = Pt(11) #11号字

def Head2():
    H2 = doc.add_paragraph()
    H2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    H2.paragraph_format.line_spacing = 1.5 # 1.5倍行距 单倍行距 1.0
    R2 = H2.add_run("——周报")
    R2.font.bold = True
    R2.font.size = Pt(11)
    doc.save(PATH)

    doc.add_page_break()

# In[6]:
def chapter1(doc,file_description,data,fixed_txt):
    h1 = doc.add_paragraph()
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1.paragraph_format.line_spacing = Pt(24)
    r1 = h1.add_run("一、")
    r1.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r1.font.bold = True
    r1.font.size = Pt(14) #四号字体对应14磅
    add_h1(h1,'上周市场走势总结',doc,file_description)

    add_p()
    add_d('日均两市成交额',doc,file_description)
    add_r('日均两市成交额',doc,**data)
    add_r('截止最后一日两市成交额',doc,**data)
    add_pic(PATH_1 + "\周报作图\liangshi.jpg")
    add_pic(PATH_1 + "\周报作图\TTransactionAmount.jpg")
    add_p()
    add_d('日均指数换手率',doc,file_description)
    add_r('日均指数换手率',doc,**data)
    add_pic(PATH_1 + r"\周报作图\turnover.jpg")
    add_p()
    add_d('申万指数收益率',doc,file_description)
    add_r('申万指数收益率',doc,**data)
    add_pic(PATH_1 + "\周报作图\SW_index.jpg")

    add_pg('极差波动率估计量',doc,**fixed_txt)
    add_pic(PATH_1 + '\素材\极差波动率.png') #这里的路径使用时需要更改

    add_p()
    add_d('截止最后一日各指数波动率',doc,file_description)
    add_r('截止最后一日各指数波动率',doc,**data)
    add_pic(PATH_1 + "\周报作图\指数日内波动.jpg")

    add_p()
    add_r('个股趋同度',doc,**fixed_txt)
    add_d('日均指数个股趋同度',doc,file_description)
    add_r('日均指数个股趋同度',doc,**data)
    add_pic(PATH_1 + "\周报作图\个股趋同度.jpg")

    add_p()
    add_r('行情集中度',doc,**fixed_txt)
    add_d('日均指数行情集中度',doc,file_description)
    add_r('日均指数行情集中度',doc,**data)
    add_pic(PATH_1 + "\周报作图\行情集中度.jpg")

    add_p()
    # add_r('RSI30D',doc,**fixed_txt)
    add_d('指数RSI30D',doc,file_description)
    add_r('指数RSI30D',doc,**data)
    add_pic(PATH_1 + "\周报作图\RSI30D.jpg")

    add_p()
    add_r('涨跌停的股票',doc,**fixed_txt)
    add_r('日均涨跌停数量',doc,**data)
    add_pic(PATH_1 + "\周报作图\涨跌停.jpg")
    add_p()
    add_r('指数成分波动率',doc,**fixed_txt)
    date_period()
    add_pg('指数成分波动率例子',doc,**fixed_txt)
    add_pic(PATH_1 + "\周报作图\指数成分波动率012.jpg")
    add_pic(PATH_1 + "\周报作图\指数成分波动率210.jpg")
    excel_pic(0,path_pic,sheet_name)
    add_pic(path_pic[0])
    excel_pic(1,path_pic,sheet_name)
    add_pic(path_pic[1])
    excel_pic(2,path_pic,sheet_name)
    add_pic(path_pic[2])

    add_p()
    add_r('指数成分波动率期限分布',doc,**fixed_txt)
    date_period()
    add_pg('指数成分波动率期限分布例子',doc,**fixed_txt)
    add_pic(PATH_1 + "\周报作图\指数成分波动率期限分布012.jpg")
    add_pic(PATH_1 + "\周报作图\指数成分波动率期限分布210.jpg")
    excel_pic(3,path_pic,sheet_name)
    add_pic(path_pic[3])
    excel_pic(4,path_pic,sheet_name)
    add_pic(path_pic[4])
    excel_pic(5,path_pic,sheet_name)
    add_pic(path_pic[5])

    add_p()
    add_r('指数成分收益率',doc,**fixed_txt)
    date_period()
    add_pic(PATH_1 + "\周报作图\指数成分收益率.jpg")

    add_pg('龙虎榜',doc,**fixed_txt)
    add_pic(PATH_1 + "\周报作图\龙虎榜.jpg")
    add_pg('日均龙虎榜上榜数量',doc,**data)

    add_pg('中信股票量化策略指数系列',doc,**fixed_txt)

    add_p()
    add_r('重要股东增持因子指数',doc,**data)
    add_pg('重要股东',doc,**fixed_txt)
    add_pic(PATH_1 + "\周报作图\IMshareholder.jpg")

    add_pg('融资融券余额',doc,**data)
    add_pic(PATH_1 + "\周报作图\liangshi_Securities_margin_trading.jpg")
    add_pic(PATH_1 + "\周报作图\Total_Securities_margin_trading.jpg")

    add_pg('可转债A股流动性指数',doc,**fixed_txt)
    add_pg('可转债市场流动性指标',doc,**data)
    add_pg('A股市场流动性指标',doc,**data)
    add_pic(PATH_1 + "\周报作图\stk_liquidity_index.jpg")
    add_pic(PATH_1 + "\周报作图\cb_liquidity_index.jpg")

    doc.add_paragraph("ETF 上周的月净申购数据和周净申购数据如下（亿元）： ")
    doc.add_paragraph("月净申购数据：")
    excel_pic(6,path_pic,sheet_name)
    add_pic(path_pic[6])
    excel_pic(7,path_pic,sheet_name)
    add_pic(path_pic[7])
    excel_pic(8,path_pic,sheet_name)
    add_pic(path_pic[8])
    doc.add_paragraph("周净申购数据：")
    excel_pic(9,path_pic,sheet_name)
    add_pic(path_pic[9])
    excel_pic(10,path_pic,sheet_name)
    add_pic(path_pic[10])
    excel_pic(11,path_pic,sheet_name)
    add_pic(path_pic[11])
    doc.add_paragraph("ETF 规模及份额（亿元）：")
    excel_pic(12,path_pic,sheet_name)
    add_pic(path_pic[12])
    excel_pic(13,path_pic,sheet_name)
    add_pic(path_pic[13])
    excel_pic(14,path_pic,sheet_name)
    add_pic(path_pic[14])
    excel_pic(15,path_pic,sheet_name)
    add_pic(path_pic[15])
    doc.add_page_break()

# In[8]:
def chapter2(doc):
    h2 = doc.add_paragraph()
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = Pt(24)
    r2 = h2.add_run("二、期货行情")
    r2.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r2.font.bold = True
    r2.font.size = Pt(14) #四号字体对应14磅

    p2 = doc.add_paragraph()
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    
    add_pic2(p2,PATH_1 + r"\周报作图\future_AL.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_B.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_BU.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_C.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_CF.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_CS.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_CU.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_HC.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_I.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_IC.CFE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_IF.CFE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_IH.CFE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_J.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_JM.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_M.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_MA.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_NI.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_OI.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_PB.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_PP.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_RB.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_RM.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_RU.SHF.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_SC.INE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_SF.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_SM.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_SR.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_TA.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_V.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_Y.DCE.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_ZC.CZC.jpg")
    add_pic2(p2,PATH_1 + r"\周报作图\future_ZN.SHF.jpg")
    doc.add_page_break()

# In[10]:
def chapter3(doc,file_description,data,fixed_txt):
    h3 = doc.add_paragraph()
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.line_spacing = Pt(24)
    r3 = h3.add_run("三、行业板块数据 – 行业涨跌分化")
    r3.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r3.font.bold = True
    r3.font.size = Pt(14) #四号字体对应14磅

    add_p()
    add_d('行业收益率',doc,file_description)
    add_pic(PATH_1 + r"\周报作图\中信一级行业指数上周收益.jpg")
    add_pg('行业收益率',doc,**data)
    add_pic(PATH_1 + r"\周报作图\中信一级行业指数.jpg")
    add_pg('行业估值变化',doc,**data)
    add_pic(PATH_1 + r"\周报作图\中信一级行业相对估值.jpg")
    add_p()
    add_d('各板块相对估值变化',doc,file_description)
    add_r('板块相对估值',doc,**fixed_txt)
    add_r('各板块相对估值变化',doc,**data)
    add_pic(PATH_1 + r"\周报作图\成长价值周期相对估值.jpg")
    doc.add_page_break()
# In[12]:
def chapter4(doc,file_description,data,data_sentiment,fixed_txt):
    h4 = doc.add_paragraph()
    h4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4.paragraph_format.line_spacing = Pt(24)
    r4 = h4.add_run("四、市场情绪")
    r4.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r4.font.bold = True
    r4.font.size = Pt(14) #四号字体对应14磅

    add_p()
    add_d('A股风险溢价',doc,file_description)
    add_r('A股风险溢价',doc,**data)
    add_pg('股权风险溢价',doc,**fixed_txt)
    add_pic(PATH_1 + r"\周报作图\A_RP.jpg")
    add_pg('Sentiment指标', doc, **fixed_txt)
    add_r('截止最后一个交易日sentiment变化', doc, **data_sentiment)
    add_r('日均sentiment变化', doc, **data_sentiment)
    add_pic(PATH_1 + r"\Sentiment\sentiment_daily.jpg")

    add_pg('人民币资金面情绪指数1', doc, **fixed_txt)
    add_pg('人民币资金面情绪指数2', doc, **fixed_txt)
    add_pg('人民币资金面情绪指数3', doc, **fixed_txt)
    add_pg('人民币资金面情绪指数4', doc, **fixed_txt)
    add_p()
    add_d('日均人民币情绪指数', doc, file_description)
    add_r('日均人民币情绪指数', doc, **data_sentiment)
    add_pic(PATH_1 + r"\Sentiment\rmb_sentiment.jpg")

    add_p()
    add_r('VIX',doc,**fixed_txt)
    add_r('VIX',doc,**data)
    add_pic(PATH_1 + r"\周报作图\VIX.jpg")
    add_p()
    add_r('股指期货贴水',doc,**fixed_txt)
    add_d('期货次月合约',doc,file_description)
    add_r('期货次月合约',doc,**data)
    add_pic(PATH_1 + r"\周报作图\次月合约贴水率.jpg")
    add_pg('中国经济政策不确定性指数',doc,**fixed_txt)
    add_pic(PATH_1 + r"\周报作图\EPU.jpg")
    doc.add_page_break()
# In[14]:
def chapter5(doc,file_description,data,fixed_txt):
    h5 = doc.add_paragraph()
    h5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h5.paragraph_format.line_spacing = Pt(24)
    r5 = h5.add_run("五、资金流入")
    r5.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r5.font.bold = True
    r5.font.size = Pt(14) #四号字体对应14磅

    add_p()
    add_d('股票质押',doc,file_description)
    excel_pic(16,path_pic,sheet_name)
    add_pic(path_pic[16])
    add_p()
    add_d('北上资金',doc,file_description)
    add_r('截止最后一个交易日沪深港通流入流出',doc,**data)
    add_pic(PATH_1 + r"\周报作图\北向资金.jpg")
    add_pg('日均主板流入资金',doc,**data)
    add_pic(PATH_1 + r"\周报作图\block_cash.jpg")
    add_pg('外资流入',doc,**fixed_txt)
    add_p()
    add_d('外资净流入',doc,file_description)
    global p
    p.add_run('A股的全球配置需求回升，持续维持在较高水平。')
    add_pic(PATH_1 + r"\Sentiment\foreign_capital.jpg")
    add_pg('交易性外资流入资金',doc,**data)
    add_pic(PATH_1 + r"\周报作图\foreign_trading_captal.jpg")

    doc.add_page_break()

    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(21) #两个五号字体10.5 * 2 = 21磅
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    doc.add_paragraph("下面是分行业的外资流入监控")
    doc.save(PATH)

    p5 = doc.add_paragraph()
    p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p5.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示

    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_传媒.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_电力.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_电子.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_房地.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_纺织.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_非银.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_钢铁.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_国防.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_机械.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_基础.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_计算.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_家电.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_建材.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_建筑.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_交通.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_煤炭.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_农林.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_汽车.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_轻工.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_商贸.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_石油.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_食品.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_通信.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_消费.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_医药.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_银行.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_有色.jpg")
    add_pic2(p5,PATH_1 + r"\周报作图\foreign_captal_综合.jpg")
    doc.add_page_break()
# In[16]:
def chapter6(doc,file_description,data):
    h6 = doc.add_paragraph()
    h6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h6.paragraph_format.line_spacing = Pt(24)
    r6 = h6.add_run("六、资金流出")
    r6.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r6.font.bold = True
    r6.font.size = Pt(14) #四号字体对应14磅

    add_p()
    add_d('解禁数量及市值',doc,file_description)
    add_r('解禁数量及市值',doc,**data)
    excel_pic(17,path_pic,sheet_name)
    add_pic(path_pic[17])
    add_pg('沪港通额度统计',doc,**data)
    add_pic(PATH_1 + r"\周报作图\限售股解禁市值.jpg")
    add_p()
    add_r('深港通额度统计',doc,**data)
    add_d('日均深股通',doc,file_description)
    excel_pic(18,path_pic,sheet_name)
    add_pic(path_pic[18])
    excel_pic(19,path_pic,sheet_name)
    add_pic(path_pic[19])
    doc.add_page_break()
# In[18]:
def chapter7(doc,file_description,data,fixed_txt):
    h7 = doc.add_paragraph()
    h7.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h7.paragraph_format.line_spacing = Pt(24)
    r7 = h7.add_run("七、利率及汇率")
    r7.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r7.font.bold = True
    r7.font.size = Pt(14) #四号字体对应14磅

    add_p()
    add_d('信用利差',doc,file_description)
    add_r('信用利差',doc,**data)
    add_pic(PATH_1 + r"\周报作图\Credit_Spread.jpg")
    add_p()
    add_d('GC007_DR007',doc,file_description)
    add_r('DRCG',doc,**fixed_txt)
    add_r('GC007_DR007',doc,**data)
    add_pic(PATH_1 + r"\周报作图\GC007_DR007.jpg")
    add_pg('逆回购',doc,**fixed_txt)
    add_pg('逆回购金额',doc,**data)
    add_pic(PATH_1 + r"\周报作图\Reverse_Repo.jpg")
    add_pg('TB_termspread',doc,**data)
    add_pic(PATH_1 + r"\周报作图\TB_termspread.jpg")
    add_pg('国债到期收益率',doc,**data)
    add_pic(PATH_1 + r"\周报作图\国债到期收益率.jpg")
    add_pg('10年期企业债(AAA)收益率_10年期国开债收益率',doc,**data)
    add_pic(PATH_1 + r"\周报作图\10年期企业债(AAA)收益率_10年期国开债收益率.jpg")
    add_pg('中短期票据收益率',doc,**data)
    add_pic(PATH_1 + r"\周报作图\S_M_Bill.jpg")
    add_pg('USDCNH',doc,**data)
    add_pic(PATH_1 + r"\周报作图\USDCNH.jpg")
    add_pg('美国国债到期收益率',doc,**data)
    add_pic(PATH_1 + r"\周报作图\USCostofCarry.jpg")

# In[20]:
def chapter8(doc,file_description,fixed_txt):
    h8 = doc.add_paragraph()
    h8.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h8.paragraph_format.line_spacing = Pt(24)
    r8 = h8.add_run("八、风格因子收益表现监测")
    r8.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r8.font.bold = True
    r8.font.size = Pt(14) #四号字体对应14磅

    add_pg('风险因子',doc,**fixed_txt)
    add_p()
    add_d('风格因子月度收益',doc,file_description)
    add_pic(PATH_1 + r"\周报作图\Barra.jpg")

# In[22]:
def chapter9(doc):
    h9 = doc.add_paragraph()
    h9.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h9.paragraph_format.line_spacing = Pt(24)
    r9 = h9.add_run("九、流动股大股东持股情况")
    r9.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r9.font.bold = True
    r9.font.size = Pt(14) #四号字体对应14磅
    excel_pic(20,path_pic,sheet_name)
    add_pic(path_pic[20])
    
# In[23]:
def chapter10(doc,fixed_txt):
    h10 = doc.add_paragraph()
    h10.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h10.paragraph_format.line_spacing = Pt(24)
    r10 = h10.add_run("十、市场方向预测")
    r10.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r10.font.bold = True
    r10.font.size = Pt(14) #四号字体对应14磅

    add_pg('市场方向预测',doc,**fixed_txt)
    excel_pic(21,path_pic,sheet_name)
    add_pic(path_pic[21])
    excel_pic(22,path_pic,sheet_name)
    add_pic(path_pic[22])

# In[25]:
def chapter11(doc,fixed_txt):
    h11 = doc.add_paragraph()
    h11.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h11.paragraph_format.line_spacing = Pt(24)
    r11 = h11.add_run("十一、相似日期检索：")
    r11.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r11.font.bold = True
    r11.font.size = Pt(14) #四号字体对应14磅

    add_pg("状态向量",doc,**fixed_txt)
    add_pg("股票市场",doc,**fixed_txt)
    add_pg("vol",doc,**fixed_txt)
    add_pg("股指期货收益率/指数收益率/集中度/趋同度：",doc,**fixed_txt)
    add_pg("大盘相对于小盘",doc,**fixed_txt)
    add_pg("商品期货市场：",doc,**fixed_txt)
    add_pg("行业行情：",doc,**fixed_txt)
    add_pg("因子收益率：",doc,**fixed_txt)
    doc.save(PATH)

    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(21) #两个五号字体10.5 * 2 = 21磅
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    add_pg('状态向量规律',doc,**fixed_txt)
    doc.add_paragraph("预测搜寻自2012年02月01日至预测日，是否存在相同状态向量的日期：")
    doc.add_paragraph("（1）以下一个月预测值为条件的匹配结果作图如下（2021-12-17）：")
    doc.add_paragraph("  股票市场：")
    add_pic(PATH_1 + r"\相似日期\stock_market.jpg")
    doc.add_paragraph("   商品期货市场：")
    add_pic(PATH_1 + r"\相似日期\commodity_market.jpg")
    doc.add_paragraph("  行业行情：")
    add_pic(PATH_1 + r"\相似日期\industries.jpg")
    doc.add_paragraph("  因子收益率：")
    add_pic(PATH_1 + r"\相似日期\factor_return.jpg")
    doc.add_paragraph("(2)对历史相似日期与预测相似日期取交集，得到的日期序列可以视为与该日期接下来一个月最相似的日期集合。交集日期具体为：")
    doc.add_paragraph("股票市场：")
    read_union_stock_market()
    doc.add_paragraph("期货市场：")
    read_union_commodity_market()
    doc.add_paragraph("行业分布：")
    read_union_industries()
    doc.add_paragraph("因子收益：")
    read_union_factor_return()

# In[27]:
def chapter12(doc):
    h12 = doc.add_paragraph()
    h12.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h12.paragraph_format.line_spacing = Pt(24)
    r12 = h12.add_run("十二、分层方向预测：")
    r12.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r12.font.bold = True
    r12.font.size = Pt(14) #四号字体对应14磅

    p = doc.add_paragraph("采用过去1年预测AUC最高的行业集中度(记为JZ)或个股趋同度(记为QT)的预测状态作为筛选条件，在满足筛选条件下，决定当天是否采用上述预测的信号。")
    p = doc.add_paragraph("筛选条件包括以下四个情境：")
    p = doc.add_paragraph("Case 1:JZ/QT预测为1")
    p = doc.add_paragraph("Case 2：JZ/QT预测为-1")
    p = doc.add_paragraph("Case 3：JZ/QT预测与待预测对象预测同向")
    p = doc.add_paragraph("Case 4：JZ/QT预测与待预测对象反向")
    p = doc.add_paragraph("其中，Case 1和Case 2是对立事件，Case 3和Case 4是对立事件。")
    p = doc.add_paragraph("无分层结果表明以历史数据考评的最优情景没有发生，前一周的月频分层情况如下：")
    p = doc.add_paragraph("以历史数据考评的最优情景没有发生,自2019年至今，分层预测的正确率如下：")
    excel_pic(23,path_pic,sheet_name)
    add_excel_pic(path_pic[23])

# In[29]:
def chapter13(doc):
    h13 = doc.add_paragraph()
    h13.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h13.paragraph_format.line_spacing = Pt(24)
    r13 = h13.add_run("十三、预测特征及特征值重要性")
    r13.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r13.font.bold = True
    r13.font.size = Pt(14) #四号字体对应14磅
    doc.save(PATH)
    r13_2 = h13.add_run("（只列示选用的前三个）")
    r13_2.font.color.rgb = RGBColor(31,73,125) # 蓝色
    r13_2.font.bold = True
    # r13_2.font.size = Pt(10.5)
    #原先的周报前后字体不一致，这里按照正文全部五号字体进行统一
    excel_pic(24,path_pic,sheet_name)
    add_pic(path_pic[24])

    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5 # 若设置为Pt，用绝对值Pt（）磅来表示
    r = p.add_run("以下列示预测的热图：")
    r.font.color.rgb  =  RGBColor(31,73,125) # 蓝色
    r.font.bold = True
    # r13.font.size = Pt(10.5)

    add_pic(PATH_1 + r"\热图\vol_m.png")
    add_pic(PATH_1 + r"\热图\Barra_return_m.png")
    add_pic(PATH_1 + r"\热图\Ashare_index_m.png")
    add_pic(PATH_1 + r"\热图\Zx_index_m.png")
    add_pic(PATH_1 + r"\热图\Ashare_feature_m.png")
    add_pic(PATH_1 + r"\热图\future_list_m.png")

if __name__ == '__main__':
    file_import()
    doc_create()
    PATH  = doc_create()
    # doc = docx.Document(r'.\素材\base.docx')
    # 后面写入内容之后再doc=Document(PATH)
    # 否则直接doc = Document(PATH)会报错，因为是空白文档
    file_import()
    data = file_import()[0]
    data_sentiment = file_import()[1]
    file_description = file_import()[2]
    fixed_txt = file_import()[3]
    doc_style()
    path_pic,sheet_name = path_pic()
#     excel_pic()
    Head1()
    Head2()
    chapter1(doc,file_description,data,fixed_txt)
    chapter2(doc)
    chapter3(doc,file_description,data,fixed_txt)
    chapter4(doc,file_description,data,data_sentiment,fixed_txt)
    chapter5(doc,file_description,data,fixed_txt)
    chapter6(doc,file_description,data)
    chapter7(doc,file_description,data,fixed_txt)
    chapter8(doc,file_description,fixed_txt)
    chapter9(doc)
    chapter10(doc,fixed_txt)
    chapter11(doc,fixed_txt)
    chapter12(doc)
    chapter13(doc)
    print('OP周报生成！！！')

